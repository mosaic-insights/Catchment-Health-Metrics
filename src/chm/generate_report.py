from __future__ import annotations

# ── stdlib
import os
import glob
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

# ── third-party
import numpy as np
import pandas as pd
import geopandas as gpd
import matplotlib.pyplot as plt
from matplotlib.image import imread
from matplotlib.ticker import MaxNLocator, ScalarFormatter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ===================== configuration =====================
@dataclass
class ReportConfig:
    chm_workspace: str
    catchment_path: str

    # NEW unified daily csv (preferred)
    historical_daily_filename: str = "{catchment}_historical_daily.csv"

    # Legacy fallbacks (kept for backward-compat)
    awap_filename: str = "{catchment}_AWAP_precip_temp.csv"
    awral_filename: str = "{catchment}_AWRAL_daily.csv"

    # subfolders
    catch_plots_dir: str = "Catchment Plots and Maps"
    catch_datasets_dir: str = "Catchment Datasets"
    sites_datasets_dir: str = "Sites Datasets"
    sites_plots_dir: str = "Sites Plots and Maps"

    # SPI/SRI products
    spi12_monthly_csv: str = "{catchment}_SPI12_monthly.csv"
    spi12_annual_csv: str = "{catchment}_SPI12_annual.csv"
    sri12_monthly_csv: str = "{catchment}_SRI12_monthly.csv"
    sri12_annual_csv: str = "{catchment}_SRI12_annual.csv"
    spi12_png: str = "{catchment}_SPI12.png"
    sri12_png: str = "{catchment}_SRI12.png"

    # sizes
    hydro_ts_size: Tuple[int, int] = (10, 15)
    bar_ts_size: Tuple[int, int] = (7, 4)
    default_dpi: int = 300

    # sites gpkg fallback
    all_sites_gpkg_name: str = "{catchment} Sites Data.gpkg"

    # land/erosion panel year window
    le_year_min: int = 2010
    le_year_max: int = 2024

    # dependency flags
    enable_spi_sri: bool = True

    # flexible name patterns
    hydro_ts_patterns: List[str] = field(default_factory=lambda: [
        "*Hydroclimate*Time*Series*Annual*.png",
        "*Hydroclimate*Annual*.png",
        "*Hydro*Annual*.png",
    ])
    monitoring_plot_patterns: List[str] = field(default_factory=lambda: [
        "Site_*_Monitoring*TimeSeries*.png",
        "*Monitoring*TimeSeries*.png",
        "*Monitoring*.png",
        "*TimeSeries*.png",
    ])
    basemap_patterns: List[str] = field(default_factory=lambda: [
        "*Catchment*Stream*Network*.png", "*Basemap*.png", "*Catchment*Map*.png"
    ])
    landuse_patterns: List[str] = field(default_factory=lambda: [
        "*Landuse*Grouped*.png", "*Landuse*.png", "*Land*Use*.png"
    ])

    # DEA Land Cover (two-panel image output name)
    dea_two_panel_name: str = "{catchment}_DEA_LandCover_FirstLast.png"

    # Canonical LU time-series images (insert after DEA two-panel)
    lu_pct_overtime_name: str = "Catchment_LU_Percentages_OverTime.png"
    lu_pct_pointchange_name: str = "Catchment_LU_PercentagePointChange.png"


# ===================== small utils =====================
def _log(status: str, message: str) -> None:
    print(f"[{status}] {message}")


def _ensure_dirs(paths: Iterable[Path]) -> None:
    for p in paths:
        p.mkdir(parents=True, exist_ok=True)


def _fmt_float(x: float) -> str:
    if isinstance(x, (float, np.floating)):
        return "" if not np.isfinite(x) else f"{x:.3g}"
    return str(x)


def _first_existing(paths: Iterable[Path]) -> Optional[Path]:
    for p in paths:
        if p and p.exists():
            return p
    return None


def _catchment_name(path: Path) -> str:
    return path.stem.replace("_", " ")


def _insert_figure(doc: Document, fig_path: Path, caption: str, width_in: float = 6.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(fig_path), width=Inches(width_in))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _safe_num(v) -> float:
    try:
        return float(v)
    except Exception:
        return np.nan


def _glob_first(folder: Path, patterns: List[str]) -> Optional[Path]:
    for pat in patterns:
        hits = sorted(folder.glob(pat))
        if hits:
            return hits[0]
    return None


# ===================== I/O: load unified & legacy =====================
def _load_historical(cfg: ReportConfig, base: Path, catchment_name: str) -> Tuple[pd.DataFrame, str]:
    datasets = base / cfg.catch_datasets_dir
    hydro = datasets / "Hydroclimate"
    hist_dir = hydro / "Historical data"

    candidates = [
        hist_dir / cfg.historical_daily_filename.format(catchment=catchment_name),
        datasets / cfg.historical_daily_filename.format(catchment=catchment_name),
    ]
    hist_path = _first_existing(candidates)

    if hist_path:
        try:
            df = pd.read_csv(hist_path)
            dcol = next((c for c in df.columns if c.lower().startswith("date")), "Date")
            df[dcol] = pd.to_datetime(df[dcol], dayfirst=True, errors="coerce")
            df = df.dropna(subset=[dcol]).sort_values(dcol)
            for c in ["Precipitation", "Min Temperature", "Max Temperature",
                      "Runoff", "Actual ET", "Upper Soil Moisture", "Deeper Soil Moisture"]:
                if c in df.columns:
                    df[c] = pd.to_numeric(df[c], errors="coerce")
            _log("OK", f"Historical daily CSV loaded: {hist_path}")
            return df, dcol
        except Exception as e:
            _log("WARN", f"Failed to read unified daily CSV at {hist_path}: {e}")

    awap_candidates = [
        datasets / cfg.awap_filename.format(catchment=catchment_name),
        hydro / "Historical AWAP" / cfg.awap_filename.format(catchment=catchment_name),
    ]
    awral_candidates = [
        datasets / cfg.awral_filename.format(catchment=catchment_name),
        hydro / "Historical AWRAL" / cfg.awral_filename.format(catchment=catchment_name),
    ]
    awap_path = _first_existing(awap_candidates)
    awral_path = _first_existing(awral_candidates)

    awap = pd.DataFrame(); awral = pd.DataFrame()
    d_awap = "Date"; d_awral = "Date"

    if awap_path:
        try:
            awap = pd.read_csv(awap_path)
            d_awap = next((c for c in awap.columns if c.lower().startswith("date")), "Date")
            awap[d_awap] = pd.to_datetime(awap[d_awap], dayfirst=True, errors="coerce")
            awap = awap.dropna(subset=[d_awap]).sort_values(d_awap)
            for c in ["Precipitation", "Min Temperature", "Max Temperature"]:
                if c in awap.columns:
                    awap[c] = pd.to_numeric(awap[c], errors="coerce")
            _log("OK", f"Legacy AWAP CSV loaded: {awap_path}")
        except Exception as e:
            _log("WARN", f"AWAP read failed: {e}")
            awap = pd.DataFrame()
    else:
        _log("WARN", "AWAP CSV not found for legacy fallback.")

    if awral_path:
        try:
            awral = pd.read_csv(awral_path)
            d_awral = next((c for c in awral.columns if c.lower().startswith("date")), "Date")
            awral[d_awral] = pd.to_datetime(awral[d_awral], dayfirst=True, errors="coerce")
            awral = awral.dropna(subset=[d_awral]).sort_values(d_awral)
            for c in ["Runoff", "Actual ET", "Upper Soil Moisture", "Deeper Soil Moisture"]:
                if c in awral.columns:
                    awral[c] = pd.to_numeric(awral[c], errors="coerce")
            _log("OK", f"Legacy AWRAL CSV loaded: {awral_path}")
        except Exception as e:
            _log("WARN", f"AWRAL read failed: {e}")
            awral = pd.DataFrame()
    else:
        _log("WARN", "AWRAL CSV not found for legacy fallback.")

    if awap.empty and awral.empty:
        _log("WARN", "No historical hydroclimate CSVs could be loaded.")
        return pd.DataFrame(), "Date"

    if not awap.empty and not awral.empty:
        dcol = "Date"
        awap2 = awap.rename(columns={d_awap: dcol})
        awral2 = awral.rename(columns={d_awral: dcol})
        merged = pd.merge(awap2, awral2, on=dcol, how="outer").sort_values(dcol)
        _log("OK", "Legacy AWAP and AWRAL CSVs merged.")
        return merged, dcol
    elif not awap.empty:
        return awap, d_awap
    else:
        return awral, d_awral


# ===================== SPI/SRI =====================
def _plot_spi_bar(df_year_mean: pd.DataFrame, col: str, title: str, ylabel: str, outfile_png: Path, size=(7, 4), dpi=300):
    if df_year_mean is None or df_year_mean.empty:
        return
    fig, ax = plt.subplots(figsize=size)

    vals = df_year_mean[col].values
    years = df_year_mean["Year"].values
    colors = ["deepskyblue" if v >= 1 else "tomato" if v <= -1 else "lime" for v in vals]
    ax.bar(years, vals, color=colors)

    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.set_xlabel("Year", fontsize=10, fontweight="bold")
    ax.set_ylabel(ylabel, fontsize=10, fontweight="bold")
    ax.tick_params(axis="both", labelsize=10)
    ax.grid(True, linestyle="--", alpha=0.6)

    if len(years):
        ax.set_xlim(min(years) - 1, max(years) + 1)

    ax.axhline(0, color="black", linestyle="--", linewidth=1)
    ax.axhline(1, color="blue", linestyle="--", linewidth=1)
    ax.axhline(-1, color="red", linestyle="--", linewidth=1)

    fig.tight_layout()
    fig.savefig(outfile_png, dpi=dpi)
    plt.close(fig)


def _compute_spi_sri_from_daily(
    cfg: ReportConfig,
    base: Path,
    catchment_name: str,
    daily: pd.DataFrame,
    dcol: str,
) -> Tuple[Optional[Path], Optional[Path]]:
    plots = base / cfg.catch_plots_dir
    datasets = base / cfg.catch_datasets_dir
    hydro = datasets / "Hydroclimate"
    hist_dir = hydro / "Historical data"
    _ensure_dirs([plots, datasets])

    spi_png = plots / cfg.spi12_png.format(catchment=catchment_name)
    sri_png = plots / cfg.sri12_png.format(catchment=catchment_name)

    if not cfg.enable_spi_sri:
        _log("INFO", "SPI/SRI disabled in config.")
        return (spi_png if spi_png.exists() else None, sri_png if sri_png.exists() else None)

    try:
        from standard_precip import spi
        spi_calc = spi.SPI()

        if not daily.empty and ("Precipitation" in daily.columns):
            precip_m = (
                daily.set_index(dcol)["Precipitation"]
                .resample("ME").sum().rename("Precipitation").to_frame().reset_index()
                .rename(columns={dcol: "Datetime"})
            )
            spi12 = spi_calc.calculate(
                df=precip_m, date_col="Datetime", precip_cols=["Precipitation"],
                freq="M", scale=12, fit_type="lmom", dist_type="gam"
            )
            spi_col = spi12.columns[2]
            spi12 = spi12.rename(columns={spi_col: "SPI12"})
            spi12.to_csv(hydro / cfg.spi12_monthly_csv.format(catchment=catchment_name), index=False)
            spi12["Year"] = spi12["Datetime"].dt.year
            spi_year = spi12.groupby("Year", as_index=False)["SPI12"].mean()
            spi_year.to_csv(hydro / cfg.spi12_annual_csv.format(catchment=catchment_name), index=False)
            _plot_spi_bar(
                spi_year, "SPI12",
                f"{catchment_name} – SPI12 (Precipitation)", "SPI12",
                spi_png, size=cfg.bar_ts_size, dpi=cfg.default_dpi
            )
            _log("OK", f"SPI12 figure generated: {spi_png}")
        else:
            _log("WARN", "SPI12 skipped: 'Precipitation' not present in daily table.")

        if not daily.empty and ("Runoff" in daily.columns):
            runoff_m = (
                daily.set_index(dcol)["Runoff"]
                .resample("ME").sum().rename("Runoff").to_frame().reset_index()
                .rename(columns={dcol: "Datetime"})
            )
            sri12 = spi_calc.calculate(
                df=runoff_m, date_col="Datetime", precip_cols=["Runoff"],
                freq="M", scale=12, fit_type="lmom", dist_type="gam"
            )
            sri_col = sri12.columns[2]
            sri12 = sri12.rename(columns={sri_col: "SRI12"})
            sri12.to_csv(hydro / cfg.sri12_monthly_csv.format(catchment=catchment_name), index=False)
            sri12["Year"] = sri12["Datetime"].dt.year
            sri_year = sri12.groupby("Year", as_index=False)["SRI12"].mean()
            sri_year.to_csv(hydro / cfg.sri12_annual_csv.format(catchment=catchment_name), index=False)
            _plot_spi_bar(
                sri_year, "SRI12",
                f"{catchment_name} – SRI12 (Runoff)", "SRI12",
                sri_png, size=cfg.bar_ts_size, dpi=cfg.default_dpi
            )
            _log("OK", f"SRI12 figure generated: {sri_png}")
        else:
            _log("WARN", "SRI12 skipped: 'Runoff' not present in daily table.")
    except Exception as e:
        _log("WARN", f"SPI/SRI step skipped: {e}")

    return (spi_png if spi_png.exists() else None, sri_png if sri_png.exists() else None)


# ===================== annual aggregation (with Mean Temperature) =====================
def _annual_from_daily(daily: pd.DataFrame, dcol: str) -> pd.DataFrame:
    if daily.empty or dcol not in daily.columns:
        return pd.DataFrame()

    df = daily.copy()
    df[dcol] = pd.to_datetime(df[dcol], errors="coerce")
    df = df.dropna(subset=[dcol]).set_index(dcol)

    if ("Min Temperature" in df.columns) and ("Max Temperature" in df.columns):
        df["Mean Temperature"] = 0.5 * (
            pd.to_numeric(df["Min Temperature"], errors="coerce")
            + pd.to_numeric(df["Max Temperature"], errors="coerce")
        )

    sum_vars = [c for c in ["Precipitation", "Runoff", "Actual ET"] if c in df.columns]
    mean_vars = [c for c in ["Min Temperature", "Max Temperature", "Mean Temperature",
                             "Upper Soil Moisture", "Deeper Soil Moisture"] if c in df.columns]
    others = [c for c in df.columns if (c not in sum_vars and c not in mean_vars)]

    pieces = []
    if sum_vars:
        pieces.append(df[sum_vars].resample("YE").sum(min_count=1))
    if mean_vars:
        pieces.append(df[mean_vars].resample("YE").mean())
    if others:
        pieces.append(df[others].resample("YE").mean())

    if not pieces:
        return pd.DataFrame()

    annual = pd.concat(pieces, axis=1)
    annual.index = annual.index.to_period("Y").to_timestamp("Y")
    return annual


# ===================== hydroclimate plot =====================
def _hydro_ts_plot(cfg: ReportConfig, base: Path, catchment_name: str, annual: pd.DataFrame) -> Optional[Path]:
    plots_dir = base / cfg.catch_plots_dir

    reuse = _glob_first(plots_dir, cfg.hydro_ts_patterns)
    if reuse:
        _log("INFO", f"Reusing existing hydroclimate plot: {reuse}")
        return reuse

    if annual is None or annual.empty:
        _log("WARN", "Hydroclimate annual dataframe is empty; no hydro plot generated.")
        return None

    panel_order = ["Precipitation", "Mean Temperature", "Runoff", "Actual ET", "Upper Soil Moisture", "Deeper Soil Moisture"]
    series = {k: annual[k].dropna() for k in panel_order if k in annual.columns and annual[k].dropna().size > 0}
    if not series:
        _log("WARN", "No valid hydroclimate annual series found for plotting.")
        return None

    def _to_year_xy(s):
        years = s.index.to_period("Y").year if hasattr(s.index, "to_period") else s.index.year
        return years.astype(int), s.values

    fig, axes = plt.subplots(nrows=len(panel_order), ncols=1, figsize=cfg.hydro_ts_size, sharex=True)
    label_kwargs = dict(fontsize=10, fontweight="bold")
    line_kwargs = dict(color="gray", linewidth=1, marker="o", markersize=3,
                       markerfacecolor="black", markeredgecolor="black")

    xmin = None; xmax = None
    for name, s in series.items():
        x, _ = _to_year_xy(s)
        if x.size:
            xmin = int(x.min()) if xmin is None else min(xmin, int(x.min()))
            xmax = int(x.max()) if xmax is None else max(xmax, int(x.max()))

    for ax, name in zip(axes, panel_order):
        ax.grid(True, linestyle="--", alpha=0.6)
        ax.set_ylabel(name, **label_kwargs)
        ax.tick_params(axis="both", labelsize=10)
        if name not in series:
            ax.text(0.5, 0.5, "Data not available", ha="center", va="center", transform=ax.transAxes)
            continue
        s = series[name]
        x, y = _to_year_xy(s)
        ax.plot(x, y, **line_kwargs)
        if np.isfinite(y).sum() >= 1:
            mean_val = np.nanmean(y)
            ax.axhline(mean_val, linestyle="--", linewidth=1, color="gray")
            ax.text(0.01, 0.90, f"Mean = {mean_val:.3g}", transform=ax.transAxes,
                    ha="left", va="center", fontsize=10, fontweight="bold")

    axes[-1].set_xlabel("Year", **label_kwargs)
    if xmin is not None and xmax is not None:
        axes[0].set_xlim([xmin - 0.5, xmax + 0.5])

    out_png = plots_dir / f"{catchment_name}_Hydroclimate_TimeSeries_Annual.png"
    out_png.parent.mkdir(parents=True, exist_ok=True)
    fig.suptitle(f"{catchment_name} – Hydroclimate Time Series (Annual)", fontsize=12, fontweight="bold")
    fig.tight_layout(rect=[0, 0.01, 1, 0.97])
    fig.savefig(out_png, dpi=cfg.default_dpi)
    plt.close(fig)
    _log("OK", f"Hydroclimate time series plot generated: {out_png}")
    return out_png


# ===================== DEA landcover helpers =====================
def _find_dea_first_last(catch_plots: Path) -> Optional[Tuple[Tuple[Path,int], Tuple[Path,int]]]:
    hits = []
    for ext in ("png", "jpg", "jpeg"):
        hits.extend(catch_plots.glob(f"Catchment_DEA_LandCover_*.{ext}"))
    years = []
    for p in hits:
        m = re.search(r"Catchment_DEA_LandCover_(\d{4})", p.name)
        if m:
            years.append((p, int(m.group(1))))
    if not years:
        return None
    years_sorted = sorted(years, key=lambda t: t[1])
    first = years_sorted[0]
    last  = years_sorted[-1]
    return first, last


def _compose_dea_two_panel(cfg: ReportConfig, catchment_name: str, catch_plots: Path) -> Optional[Path]:
    pair = _find_dea_first_last(catch_plots)
    if not pair:
        _log("WARN", "DEA first/last year images not found for two-panel composition.")
        return None
    (p_first, y_first), (p_last, y_last) = pair

    try:
        img1 = imread(str(p_first))
        img2 = imread(str(p_last))

        fig, axes = plt.subplots(
            nrows=1,
            ncols=2,
            figsize=(10.0, 5.4),
            constrained_layout=False,
        )

        panels = [(axes[0], img1, y_first), (axes[1], img2, y_last)]
        for ax, img, y in panels:
            ax.imshow(img)
            ax.axis("off")
            ax.margins(x=0, y=0)

        fig.subplots_adjust(
            left=0.02,
            right=0.99,
            bottom=0.03,
            top=0.92,
            wspace=0.001
        )

        out_path = catch_plots / cfg.dea_two_panel_name.format(catchment=catchment_name)
        fig.savefig(out_path, dpi=cfg.default_dpi, bbox_inches="tight", pad_inches=0.03)
        plt.close(fig)
        _log("OK", f"DEA two-panel image composed: {out_path}")
        return out_path

    except Exception as e:
        _log("WARN", f"Could not compose DEA two-panel: {e}")
        return None


def _add_csv_table(doc: Document, csv_path: Path, round_to: int = 3) -> None:
    try:
        df = pd.read_csv(csv_path)
        if df.empty:
            doc.add_paragraph(f"[{csv_path.name} is empty.]")
            return

        disp = df.copy()
        for c in disp.columns:
            if pd.api.types.is_numeric_dtype(disp[c]):
                disp[c] = disp[c].round(round_to)

        tbl = doc.add_table(rows=1, cols=len(disp.columns))
        hdr = tbl.rows[0].cells
        for i, c in enumerate(disp.columns):
            hdr[i].text = str(c)

        for _, r in disp.iterrows():
            cells = tbl.add_row().cells
            for i, c in enumerate(disp.columns):
                val = r[c]
                if pd.isna(val):
                    cells[i].text = ""
                else:
                    cells[i].text = str(val)
    except Exception as e:
        doc.add_paragraph(f"[Could not add table from {csv_path.name}: {e}]")


# ===================== land/erosion helpers for sites =====================
def _extract_year_series(row: pd.Series, prefix: str, suffix: str, y0: int, y1: int) -> pd.DataFrame:
    out = []
    for y in range(y0, y1 + 1):
        col = f"{prefix}{y}{suffix}"
        if col in row.index:
            out.append((y, _safe_num(row[col])))
    return pd.DataFrame(out, columns=["Year", "Value"]).sort_values("Year")


def _site_le_panel(cfg: ReportConfig, row: pd.Series, site_id, out_dir: Path) -> Optional[Path]:
    ndvi = _extract_year_series(row, "NDVI_", " (mean)", cfg.le_year_min, cfg.le_year_max)
    cf   = _extract_year_series(row, "C_Factor_", " (mean)", cfg.le_year_min, cfg.le_year_max)
    rusl = _extract_year_series(row, "RUSLE total ", " (t/yr)", cfg.le_year_min, cfg.le_year_max)
    sdr  = _extract_year_series(row, "SDR-RUSLE total ", " (t/yr)", cfg.le_year_min, cfg.le_year_max)
    if all(df.empty for df in (ndvi, cf, rusl, sdr)):
        return None

    fig, axes = plt.subplots(nrows=4, ncols=1, figsize=(10, 12), sharex=True)
    panels = [
        ("NDVI (mean)", ndvi),
        ("C-Factor (mean)", cf),
        ("RUSLE total (t/yr)", rusl),
        ("SDR-RUSLE total (t/yr)", sdr),
    ]
    for ax, (title, df) in zip(axes, panels):
        ax.grid(True, linestyle="--", alpha=0.6)
        ax.set_ylabel(title, fontsize=10, fontweight="bold")
        ax.tick_params(axis="both", labelsize=10)
        if df.empty:
            ax.text(0.5, 0.5, "Data not available", ha="center", va="center", transform=ax.transAxes)
            continue
        ax.plot(
            df["Year"], df["Value"],
            color="gray", linewidth=1, marker="o",
            markersize=4, markerfacecolor="black", markeredgecolor="black"
        )
        if np.isfinite(df["Value"]).sum() >= 1:
            mean_val = np.nanmean(df["Value"])
            ax.axhline(mean_val, linestyle="--", linewidth=1, color="gray")
            ax.text(0.01, 0.90, f"Mean = {mean_val:.3g}", transform=ax.transAxes,
                    ha="left", va="center", fontsize=10, fontweight="bold")
        ax.set_xlim(cfg.le_year_min - 0.5, cfg.le_year_max + 0.5)

    axes[-1].set_xlabel("Year", fontsize=10, fontweight="bold")
    fig.suptitle(
        f"Site {site_id} – NDVI / C-Factor / RUSLE / SDR-RUSLE ({cfg.le_year_min}–{cfg.le_year_max})",
        fontsize=11, fontweight="bold"
    )
    fig.tight_layout(rect=[0, 0.01, 1, 0.97])

    out_dir.mkdir(parents=True, exist_ok=True)
    out_png = out_dir / f"Site_{site_id}_LandErosion_TimeSeries.png"
    fig.savefig(out_png, dpi=cfg.default_dpi)
    plt.close(fig)
    return out_png


def _site_le_summary_table(row: pd.Series, cfg: ReportConfig) -> pd.DataFrame:
    def summarize(df: pd.DataFrame) -> Dict:
        if df.empty:
            return dict(Start=np.nan, End=np.nan, Mean=np.nan, Median=np.nan, Min=np.nan, Max=np.nan, Std=np.nan)
        v = pd.to_numeric(df["Value"], errors="coerce")
        v = v[np.isfinite(v)]
        if v.empty:
            return dict(Start=df["Year"].min(), End=df["Year"].max(),
                        Mean=np.nan, Median=np.nan, Min=np.nan, Max=np.nan, Std=np.nan)
        return dict(Start=int(df["Year"].min()), End=int(df["Year"].max()),
                    Mean=float(np.nanmean(v)), Median=float(np.nanmedian(v)),
                    Min=float(np.nanmin(v)), Max=float(np.nanmax(v)),
                    Std=float(np.nanstd(v, ddof=1)) if v.size >= 2 else np.nan)

    ndvi = summarize(_extract_year_series(row, "NDVI_", " (mean)", cfg.le_year_min, cfg.le_year_max))
    cf   = summarize(_extract_year_series(row, "C_Factor_", " (mean)", cfg.le_year_min, cfg.le_year_max))
    rusl = summarize(_extract_year_series(row, "RUSLE total ", " (t/yr)", cfg.le_year_min, cfg.le_year_max))
    sdr  = summarize(_extract_year_series(row, "SDR-RUSLE total ", " (t/yr)", cfg.le_year_min, cfg.le_year_max))

    rows = [
        dict(Variable="NDVI (mean)", **ndvi),
        dict(Variable="C-Factor (mean)", **cf),
        dict(Variable="RUSLE total (t/yr)", **rusl),
        dict(Variable="SDR-RUSLE total (t/yr)", **sdr),
    ]
    return pd.DataFrame(rows)


# ===================== Riparian NDVI helper functions =====================
def _find_riparian_first_last(catch_plots: Path) -> Optional[Tuple[Tuple[Path, int], Tuple[Path, int]]]:
    hits = []
    for ext in ("png", "jpg", "jpeg"):
        hits.extend(catch_plots.glob(f"Riparian_NDVI_by_StreamSegment_*.{ext}"))
    years = []
    for p in hits:
        m = re.search(r"Riparian_NDVI_by_StreamSegment_(\d{4})", p.name)
        if m:
            years.append((p, int(m.group(1))))
    if not years:
        return None
    years_sorted = sorted(years, key=lambda t: t[1])
    return years_sorted[0], years_sorted[-1]


def _compose_riparian_two_panel(cfg: ReportConfig, catchment_name: str, catch_plots: Path) -> Optional[Path]:
    pair = _find_riparian_first_last(catch_plots)
    if not pair:
        _log("WARN", "Riparian NDVI first/last year images not found.")
        return None

    (p_first, y_first), (p_last, y_last) = pair
    try:
        img1 = imread(str(p_first))
        img2 = imread(str(p_last))

        fig, axes = plt.subplots(
            nrows=1,
            ncols=2,
            figsize=(10.0, 5.4),
            constrained_layout=False,
        )

        panels = [(axes[0], img1, y_first), (axes[1], img2, y_last)]
        for ax, img, y in panels:
            ax.imshow(img)
            ax.axis("off")
            ax.margins(x=0, y=0)

        fig.subplots_adjust(
            left=0.02,
            right=0.99,
            bottom=0.03,
            top=0.92,
            wspace=0.001,
        )

        out_path = catch_plots / f"{catchment_name}_Riparian_NDVI_StreamSegment_FirstLast.png"
        fig.savefig(out_path, dpi=cfg.default_dpi, bbox_inches="tight", pad_inches=0.03)
        plt.close(fig)
        _log("OK", f"Riparian NDVI two-panel image composed: {out_path}")
        return out_path
    except Exception as e:
        _log("WARN", f"Could not compose Riparian NDVI two-panel: {e}")
        return None


# ===================== Site-level hydroclimate/SPI/SRI helpers =====================
def _load_site_daily(base: Path, cfg: ReportConfig, site_id: int) -> Tuple[pd.DataFrame, str]:
    plots_csv = base / cfg.sites_plots_dir / f"Site_{site_id}" / f"Site_{site_id}_historical_daily.csv"
    data_csv  = base / cfg.sites_datasets_dir / f"Site_{site_id}" / f"Site_{site_id}_historical_daily.csv"
    path = _first_existing([plots_csv, data_csv])
    if not path:
        _log("WARN", f"Site {site_id}: daily hydroclimate CSV not found.")
        return pd.DataFrame(), "Date"

    try:
        df = pd.read_csv(path)
        dcol = next((c for c in df.columns if c.lower().startswith("date")), "Date")
        df[dcol] = pd.to_datetime(df[dcol], dayfirst=True, errors="coerce")
        df = df.dropna(subset=[dcol]).sort_values(dcol)
        for c in ["Precipitation", "Min Temperature", "Max Temperature",
                  "Runoff", "Actual ET", "Upper Soil Moisture", "Deeper Soil Moisture"]:
            if c in df.columns:
                df[c] = pd.to_numeric(df[c], errors="coerce")
        _log("OK", f"Site {site_id}: daily hydroclimate CSV loaded.")
        return df, dcol
    except Exception as e:
        _log("WARN", f"Failed reading site daily CSV for Site {site_id}: {e}")
        return pd.DataFrame(), "Date"


def _site_hydro_ts_plot(cfg: ReportConfig, base: Path, site_id: int, annual: pd.DataFrame) -> Optional[Path]:
    site_plot_dir = base / cfg.sites_plots_dir / f"Site_{site_id}"
    site_plot_dir.mkdir(parents=True, exist_ok=True)

    if annual is None or annual.empty:
        return None

    panel_order = ["Precipitation", "Mean Temperature", "Runoff", "Actual ET", "Upper Soil Moisture", "Deeper Soil Moisture"]
    series = {k: annual[k].dropna() for k in panel_order if k in annual.columns and annual[k].dropna().size > 0}
    if not series:
        return None

    def _to_year_xy(s):
        years = s.index.to_period("Y").year if hasattr(s.index, "to_period") else s.index.year
        return years.astype(int), s.values

    fig, axes = plt.subplots(nrows=len(panel_order), ncols=1, figsize=cfg.hydro_ts_size, sharex=True)
    label_kwargs = dict(fontsize=10, fontweight="bold")
    line_kwargs = dict(color="gray", linewidth=1, marker="o", markersize=3,
                       markerfacecolor="black", markeredgecolor="black")

    xmin = None; xmax = None
    for name, s in series.items():
        x, _ = _to_year_xy(s)
        if x.size:
            xmin = int(x.min()) if xmin is None else min(xmin, int(x.min()))
            xmax = int(x.max()) if xmax is None else max(xmax, int(x.max()))

    for ax, name in zip(axes, panel_order):
        ax.grid(True, linestyle="--", alpha=0.6)
        ax.set_ylabel(name, **label_kwargs)
        ax.tick_params(axis="both", labelsize=10)
        if name not in series:
            ax.text(0.5, 0.5, "Data not available", ha="center", va="center", transform=ax.transAxes)
            continue
        s = series[name]
        x, y = _to_year_xy(s)
        ax.plot(x, y, **line_kwargs)
        if np.isfinite(y).sum() >= 1:
            mean_val = np.nanmean(y)
            ax.axhline(mean_val, linestyle="--", linewidth=1, color="gray")
            ax.text(0.01, 0.90, f"Mean = {mean_val:.3g}", transform=ax.transAxes, ha="left", va="center", fontsize=10, fontweight="bold")

    axes[-1].set_xlabel("Year", **label_kwargs)
    if xmin is not None and xmax is not None:
        axes[0].set_xlim([xmin - 0.5, xmax + 0.5])

    out_png = site_plot_dir / f"Site_{site_id}_Hydroclimate_TimeSeries_Annual.png"
    fig.suptitle(f"Site {site_id} – Hydroclimate Time Series (Annual)", fontsize=11, fontweight="bold")
    fig.tight_layout(rect=[0, 0.01, 1, 0.97])
    fig.savefig(out_png, dpi=cfg.default_dpi)
    plt.close(fig)
    return out_png


def _compute_spi_sri_for_site(
    cfg: ReportConfig,
    base: Path,
    site_id: int,
    daily: pd.DataFrame,
    dcol: str,
) -> Tuple[Optional[Path], Optional[Path], Optional[pd.DataFrame], Optional[pd.DataFrame]]:
    site_plot_dir = base / cfg.sites_plots_dir / f"Site_{site_id}"
    site_plot_dir.mkdir(parents=True, exist_ok=True)

    spi_png = site_plot_dir / f"Site_{site_id}_SPI12.png"
    sri_png = site_plot_dir / f"Site_{site_id}_SRI12.png"

    spi_year = None
    sri_year = None

    if not cfg.enable_spi_sri:
        return (spi_png if spi_png.exists() else None,
                sri_png if sri_png.exists() else None, None, None)

    try:
        from standard_precip import spi
        spi_calc = spi.SPI()

        if not daily.empty and ("Precipitation" in daily.columns):
            precip_m = (
                daily.set_index(dcol)["Precipitation"]
                .resample("ME").sum().rename("Precipitation").to_frame().reset_index()
                .rename(columns={dcol: "Datetime"}))
            spi12 = spi_calc.calculate(
                df=precip_m, date_col="Datetime", precip_cols=["Precipitation"],
                freq="M", scale=12, fit_type="lmom", dist_type="gam")
            spi_col = spi12.columns[2]
            spi12 = spi12.rename(columns={spi_col: "SPI12"})
            spi12["Year"] = spi12["Datetime"].dt.year
            spi_year = spi12.groupby("Year", as_index=False)["SPI12"].mean()

            _plot_spi_bar(
                spi_year, "SPI12",
                f"Site {site_id} – SPI12 (Precipitation)", "SPI12",
                spi_png, size=cfg.bar_ts_size, dpi=cfg.default_dpi)

        if not daily.empty and ("Runoff" in daily.columns):
            runoff_m = (
                daily.set_index(dcol)["Runoff"]
                .resample("ME").sum().rename("Runoff").to_frame().reset_index()
                .rename(columns={dcol: "Datetime"}))
            sri12 = spi_calc.calculate(
                df=runoff_m, date_col="Datetime", precip_cols=["Runoff"],
                freq="M", scale=12, fit_type="lmom", dist_type="gam")
            sri_col = sri12.columns[2]
            sri12 = sri12.rename(columns={sri_col: "SRI12"})
            sri12["Year"] = sri12["Datetime"].dt.year
            sri_year = sri12.groupby("Year", as_index=False)["SRI12"].mean()

            _plot_spi_bar(
                sri_year, "SRI12",
                f"Site {site_id} – SRI12 (Runoff)", "SRI12",
                sri_png, size=cfg.bar_ts_size, dpi=cfg.default_dpi)

    except Exception as e:
        _log("WARN", f"Site {site_id} SPI/SRI step skipped: {e}")

    return (spi_png if spi_png.exists() else None,
            sri_png if sri_png.exists() else None,
            spi_year, sri_year)


def _insert_site_plot_if_exists(doc: Document,
                                site_plot_dir: Path,
                                filename: str,
                                fig_counter: int,
                                caption: str,
                                width_in: float = 6.5
) -> int:
    p = site_plot_dir / filename
    if p.exists():
        fig_counter += 1
        _insert_figure(doc, p, f"Figure {fig_counter}. {caption}", width_in=width_in)
        _log("OK", f"Inserted plot into report: {p}")
    else:
        doc.add_paragraph(f"[{filename} not found in: {site_plot_dir}]")
        _log("WARN", f"Missing site plot for report: {p}")
    return fig_counter


SITE_PLOTS_SEQUENCE = [
    ("Site_{sid}_Riparian_NDVI_Timeseries.png",
     "Site {sid} – Riparian NDVI time series"),
    ("Site_{sid}_LU_Percentages_OverTime.png",
     "Site {sid} – Land-use percentages over time"),
    ("Site_{sid}_LU_PercentagePointChange.png",
     "Site {sid} – Land-use percentage-point change"),
    ("Site_{sid}_AUC_NDVI_SDR.png",
     "Site {sid} – AUC summary for NDVI & SDR"),
    ("Site_{sid}_AUC_Bushfire.png",
     "Site {sid} – AUC summary for Bushfire exposure"),
    ("Site_{sid}_RUSLE_SDR-RUSLE_TimeSeries.png",
     "Site {sid} – RUSLE and SDR-RUSLE (t/ha/yr)"),
    ("Site_{sid}_Road_AUC_TimeSeries.png",
     "Site {sid} – Roads: AUC time-series"),
    ("Site_{sid}_LU_Water_wetlands_AUC_TimeSeries.png","Site {sid} – AUC: Water & wetlands"),
    ("Site_{sid}_LU_Native_forests_native_grazing_AUC_TimeSeries.png","Site {sid} – AUC: Native forests & native grazing"),
    ("Site_{sid}_LU_Conservation_minimal_use_AUC_TimeSeries.png","Site {sid} – AUC: Conservation & minimal use"),
    ("Site_{sid}_LU_Plantation_forests_AUC_TimeSeries.png","Site {sid} – AUC: Plantation forests"),
    ("Site_{sid}_LU_Dryland_agriculture_AUC_TimeSeries.png","Site {sid} – AUC: Dryland agriculture"),
    ("Site_{sid}_LU_Irrigated_agriculture_AUC_TimeSeries.png","Site {sid} – AUC: Irrigated agriculture"),
    ("Site_{sid}_LU_Intensive_production_AUC_TimeSeries.png","Site {sid} – AUC: Intensive production"),
    ("Site_{sid}_LU_Urban_residential_services_AUC_TimeSeries.png", "Site {sid} – AUC: Urban, residential & services"),
    ("Site_{sid}_LU_Industry_utilities_mining_waste_AUC_TimeSeries.png","Site {sid} – AUC: Industry, utilities, mining & waste"),
    ("Site_{sid}_LU_Transport_communication_AUC_TimeSeries.png","Site {sid} – AUC: Transport & communication"),
]


def _site_location_map(
    cfg: ReportConfig,
    base: Path,
    catch_gdf: Optional[gpd.GeoDataFrame],
    site_geom,
    sid: int,
    sites_crs=None
) -> Optional[Path]:
    try:
        if site_geom is None or (hasattr(site_geom, "is_empty") and site_geom.is_empty):
            return None

        site_gdf = gpd.GeoDataFrame({"Site_id": [sid]}, geometry=[site_geom], crs=sites_crs)

        if catch_gdf is not None and catch_gdf.crs is not None:
            if site_gdf.crs is None:
                site_gdf = site_gdf.set_crs(catch_gdf.crs, allow_override=True)
            elif str(site_gdf.crs) != str(catch_gdf.crs):
                site_gdf = site_gdf.to_crs(catch_gdf.crs)

        out_dir = base / cfg.sites_plots_dir / f"Site_{sid}"
        out_dir.mkdir(parents=True, exist_ok=True)
        out_png = out_dir / f"Site_{sid}_Location.png"

        fig, ax = plt.subplots(figsize=(6.8, 6.0))

        if catch_gdf is not None and not catch_gdf.empty:
            try:
                catch_gdf.boundary.plot(ax=ax, color="black", linewidth=1.0, zorder=1)
            except Exception:
                catch_gdf.plot(ax=ax, facecolor="none", edgecolor="black", linewidth=1.0, zorder=1)

        gtype = site_gdf.geom_type.iloc[0] if not site_gdf.empty else ""
        if gtype in {"Polygon", "MultiPolygon"}:
            site_gdf.boundary.plot(ax=ax, color="crimson", linewidth=1.8, zorder=3)
            site_gdf.plot(ax=ax, facecolor="none", edgecolor="crimson", linewidth=1.8, zorder=3)
        elif gtype in {"LineString", "MultiLineString"}:
            site_gdf.plot(ax=ax, color="crimson", linewidth=2.0, zorder=3)
        else:
            site_gdf.plot(ax=ax, color="crimson", markersize=35, zorder=3)

        ax.set_title(f"Site {sid} – Location within Catchment", fontsize=11, fontweight="bold", pad=6)
        ax.set_xlabel("Longitude", fontsize=10, fontweight="bold")
        ax.set_ylabel("Latitude", fontsize=10, fontweight="bold")
        ax.xaxis.set_major_locator(MaxNLocator(nbins=4))
        ax.yaxis.set_major_locator(MaxNLocator(nbins=4))
        ax.tick_params(axis="both", labelsize=10)
        ax.grid(True, linestyle="--", alpha=0.6)

        if catch_gdf is not None and not catch_gdf.empty:
            minx, miny, maxx, maxy = catch_gdf.total_bounds
        else:
            minx, miny, maxx, maxy = site_gdf.total_bounds
            dx = (maxx - minx) or 1.0
            dy = (maxy - miny) or 1.0
            minx, maxx = minx - 0.1 * dx, maxx + 0.1 * dx
            miny, maxy = miny - 0.1 * dy, maxy + 0.1 * dy
        ax.set_xlim(minx, maxx)
        ax.set_ylim(miny, maxy)

        fig.tight_layout()
        fig.savefig(out_png, dpi=cfg.default_dpi)
        plt.close(fig)
        return out_png

    except Exception as e:
        _log("WARN", f"Site {sid} location map failed: {e}")
        return None


# ===================== document assembly =====================
def build_report(cfg: ReportConfig) -> Path:
    catch_path = Path(cfg.catchment_path)
    catchment_name = _catchment_name(catch_path)
    base = Path(cfg.chm_workspace) / catchment_name

    catch_gdf = None
    try:
        catch_gdf = gpd.read_file(cfg.catchment_path)
        _log("OK", f"Catchment boundary loaded: {cfg.catchment_path}")
    except Exception as e:
        _log("WARN", f"Could not read catchment_path: {e}")

    catch_plots = base / cfg.catch_plots_dir
    catch_datasets = base / cfg.catch_datasets_dir
    sites_datasets = base / cfg.sites_datasets_dir
    sites_plots = base / cfg.sites_plots_dir
    _ensure_dirs([base, catch_plots, catch_datasets, sites_datasets, sites_plots])

    _log("INFO", "Loading historical hydroclimate data for report...")
    daily, dcol = _load_historical(cfg, base, catchment_name)

    _log("INFO", "Computing SPI/SRI products for report...")
    spi_png, sri_png = _compute_spi_sri_from_daily(cfg, base, catchment_name, daily, dcol)

    annual = _annual_from_daily(daily, dcol)
    if annual.empty:
        _log("WARN", "Annual hydroclimate dataframe is empty.")
    else:
        _log("OK", "Annual hydroclimate dataframe prepared.")

    summary_rows = []
    for var in ["Precipitation", "Mean Temperature", "Runoff", "Actual ET",
                "Upper Soil Moisture", "Deeper Soil Moisture",
                "Min Temperature", "Max Temperature"]:
        if isinstance(annual, pd.DataFrame) and (var in annual.columns):
            s = pd.to_numeric(annual[var], errors="coerce")
            n = int(s.notna().sum())
            summary_rows.append({
                "Variable": var,
                "Mean":   float(np.nanmean(s)) if n else np.nan,
                "Median": float(np.nanmedian(s)) if n else np.nan,
                "Min":    float(np.nanmin(s)) if n else np.nan,
                "Max":    float(np.nanmax(s)) if n else np.nan,
                "Std":    float(np.nanstd(s, ddof=1)) if n >= 2 else np.nan,
            })
    summary_df = pd.DataFrame(summary_rows)

    doc = Document()
    title = doc.add_heading(level=0)
    run = title.add_run(f"This is an automated report for {catchment_name} catchment health condition")
    run.font.size = Pt(16); run.bold = True
    doc.add_heading(catchment_name, level=1)

    fig_counter = 0

    basemap = _glob_first(catch_plots, cfg.basemap_patterns)
    if basemap and basemap.exists():
        fig_counter += 1
        _insert_figure(doc, basemap, f"Figure {fig_counter}. {catchment_name} and the location of sites.", width_in=4.5)
        _log("OK", f"Added basemap figure: {basemap}")
    else:
        doc.add_paragraph(f"[Catchment basemap not found in: {catch_plots}]")
        _log("WARN", f"Catchment basemap missing in: {catch_plots}")

    lu_png = _glob_first(catch_plots, cfg.landuse_patterns)
    if lu_png and lu_png.exists():
        fig_counter += 1
        _insert_figure(doc, lu_png, f"Figure {fig_counter}. {catchment_name} land use.", width_in=6.5)
        _log("OK", f"Added land use figure: {lu_png}")
    else:
        doc.add_paragraph(f"[Grouped land-use image not found in: {catch_plots}]")
        _log("WARN", f"Grouped land-use image missing in: {catch_plots}")

    dea_two = _compose_dea_two_panel(cfg, catchment_name, catch_plots)
    if dea_two and dea_two.exists():
        fig_counter += 1
        _insert_figure(doc, dea_two,
                       f"Figure {fig_counter}. {catchment_name} DEA Land Cover.",
                       width_in=6.5)
        _log("OK", f"Added DEA two-panel figure: {dea_two}")
    else:
        doc.add_paragraph("[DEA Land Cover first/last panel not generated (files not found).]")
        _log("WARN", "DEA two-panel figure missing.")

    lu_pct_over = catch_plots / cfg.lu_pct_overtime_name
    if lu_pct_over.exists():
        fig_counter += 1
        _insert_figure(doc, lu_pct_over,
                       f"Figure {fig_counter}. {catchment_name} – Land-use percentages over time.", width_in=6.5)
        _log("OK", f"Added LU percentages over time figure: {lu_pct_over}")
    else:
        doc.add_paragraph(f"[{cfg.lu_pct_overtime_name} not found in: {catch_plots}]")
        _log("WARN", f"Missing LU percentages over time figure: {lu_pct_over}")

    lu_pct_change = catch_plots / cfg.lu_pct_pointchange_name
    if lu_pct_change.exists():
        fig_counter += 1
        _insert_figure(doc, lu_pct_change,
                       f"Figure {fig_counter}. {catchment_name} – Land-use percentage-point change.", width_in=6.5)
        _log("OK", f"Added LU percentage-point change figure: {lu_pct_change}")
    else:
        doc.add_paragraph(f"[{cfg.lu_pct_pointchange_name} not found in: {catch_plots}]")
        _log("WARN", f"Missing LU percentage-point change figure: {lu_pct_change}")

    landuse_output = catch_datasets / "Landuse" / "DEA Landcover"
    lu_pct_csv = landuse_output / "Catchment_LU_Percentages.csv"
    doc.add_paragraph("")
    doc.add_heading(f"{catchment_name} – Land-use percentages (table)", level=2)
    if lu_pct_csv.exists():
        _add_csv_table(doc, lu_pct_csv, round_to=3)
        _log("OK", f"Added LU percentages table: {lu_pct_csv}")
    else:
        doc.add_paragraph(f"[Catchment_LU_Percentages.csv not found at: {lu_pct_csv}]")
        _log("WARN", f"Missing LU percentages table: {lu_pct_csv}")

    riparian_two_panel = _compose_riparian_two_panel(cfg, catchment_name, catch_plots)
    if riparian_two_panel and riparian_two_panel.exists():
        fig_counter += 1
        _insert_figure(doc, riparian_two_panel, f"Figure {fig_counter}. {catchment_name} – Riparian NDVI by Stream Segment.", width_in=6.5)
        _log("OK", f"Added riparian two-panel figure: {riparian_two_panel}")
    else:
        doc.add_paragraph("[Riparian NDVI by Stream Segment images not found.]")
        _log("WARN", "Riparian NDVI by Stream Segment images missing.")

    riparian_timeseries = catch_plots / "Riparian_NDVI_Mean_by_StreamOrder_Timeseries.png"
    if riparian_timeseries.exists():
        fig_counter += 1
        _insert_figure(doc, riparian_timeseries, f"Figure {fig_counter}. {catchment_name} – Riparian NDVI mean by stream order over time.", width_in=6.5)
        _log("OK", f"Added riparian timeseries figure: {riparian_timeseries}")
    else:
        doc.add_paragraph("[Riparian NDVI mean by stream order timeseries missing.]")
        _log("WARN", f"Missing riparian timeseries figure: {riparian_timeseries}")

    hydro_png = _hydro_ts_plot(cfg, base, catchment_name, annual)
    if hydro_png and hydro_png.exists():
        fig_counter += 1
        _insert_figure(doc, hydro_png, f"Figure {fig_counter}. Annual hydroclimate time series for {catchment_name}.", width_in=6.5)
        _log("OK", f"Added hydroclimate timeseries figure: {hydro_png}")
    else:
        doc.add_paragraph("[Hydroclimate time series not generated.]")
        _log("WARN", "Hydroclimate time series figure missing.")

    doc.add_heading(f"Statistical summary for annual hydroclimate variables across {catchment_name}", level=2)
    cols = ["Variable", "Mean", "Median", "Min", "Max", "Std"]
    if not summary_df.empty:
        tbl = doc.add_table(rows=1, cols=len(cols))
        hdr = tbl.rows[0].cells
        for i, c in enumerate(cols): hdr[i].text = c
        for _, r in summary_df.iterrows():
            row_cells = tbl.add_row().cells
            for i, c in enumerate(cols):
                row_cells[i].text = _fmt_float(r.get(c, np.nan))
        _log("OK", "Added catchment hydroclimate summary table.")
    else:
        doc.add_paragraph("[No annual summary could be computed (missing inputs).]")
        _log("WARN", "Catchment hydroclimate summary table could not be generated.")

    if spi_png and spi_png.exists():
        fig_counter += 1
        _insert_figure(doc, spi_png, f"Figure {fig_counter}. {catchment_name} SPI-12 (annual mean of 12-month SPI).")
        _log("OK", f"Added SPI12 figure: {spi_png}")
    else:
        doc.add_paragraph("[SPI12 figure not generated.]")
        _log("WARN", "SPI12 figure missing.")

    if sri_png and sri_png.exists():
        fig_counter += 1
        _insert_figure(doc, sri_png, f"Figure {fig_counter}. {catchment_name} SRI-12 (annual mean of 12-month SRI).")
        _log("OK", f"Added SRI12 figure: {sri_png}")
    else:
        doc.add_paragraph("[SRI12 figure not generated.]")
        _log("WARN", "SRI12 figure missing.")

    rusle_sdr_means = catch_plots / f"{catchment_name}_RUSLE_SDR-RUSLE_TimeSeries.png"
    if rusle_sdr_means.exists():
        fig_counter += 1
        _insert_figure(doc, rusle_sdr_means, f"Figure {fig_counter}. {catchment_name} – RUSLE and SDR-RUSLE (t/ha/yr).", width_in=6.5)
        _log("OK", f"Added RUSLE/SDR-RUSLE figure: {rusle_sdr_means}")
    else:
        doc.add_paragraph(f"[{rusle_sdr_means.name} not found in: {catch_plots}]")
        _log("WARN", f"Missing RUSLE/SDR-RUSLE figure: {rusle_sdr_means}")

    doc.add_heading("Sites", level=1)

    site_ids: List[str] = []
    gpkg_candidates = [
        sites_datasets / cfg.all_sites_gpkg_name.format(catchment=catchment_name),
        sites_datasets / f"{catchment_name} Sites Data.gpkg",
        sites_datasets / "All Sites Data.gpkg",
    ]
    gpkg_path = _first_existing(gpkg_candidates)
    sites_gdf = None

    if gpkg_path and gpkg_path.exists():
        try:
            sites_gdf = gpd.read_file(gpkg_path)
            if "Site_id" in sites_gdf.columns:
                site_ids = sorted({str(v) for v in sites_gdf["Site_id"].dropna().unique()})
            _log("OK", f"Sites GPKG loaded for report: {gpkg_path}")
        except Exception as e:
            _log("WARN", f"Could not read sites GPKG: {e}")
            sites_gdf = None
    else:
        _log("WARN", "Sites GPKG not found; report will fall back to site folders if available.")

    if not site_ids and sites_plots.exists():
        for p in (p for p in sites_plots.iterdir() if p.is_dir()):
            if p.name.startswith("Site_"):
                site_ids.append(p.name.replace("Site_", ""))
        if site_ids:
            _log("OK", "Recovered site IDs from site plot folders.")

    site_ids = sorted(set(site_ids))

    if not site_ids:
        doc.add_paragraph("[No site IDs could be found. Site section skipped.]")
        _log("WARN", "No site IDs found. Site section skipped.")

    for sid in site_ids:
        doc.add_heading(f"Site {sid}", level=2)

        if catch_gdf is not None and sites_gdf is not None and not sites_gdf.empty and ("Site_id" in sites_gdf.columns):
            row_match = sites_gdf[sites_gdf["Site_id"].astype(str) == str(sid)]
            if not row_match.empty:
                site_row = row_match.iloc[0]
                loc_png = _site_location_map(
                    cfg, base, catch_gdf,
                    site_row.geometry,
                    sid,
                    getattr(sites_gdf, "crs", None)
                )
                if loc_png and loc_png.exists():
                    fig_counter += 1
                    _insert_figure(doc, loc_png, f"Figure {fig_counter}. Site {sid} location within {catchment_name}.", width_in=6.0)
                    _log("OK", f"Added site location map: {loc_png}")
                else:
                    doc.add_paragraph(f"[Site {sid}: location map not generated.]")
                    _log("WARN", f"Site {sid}: location map not generated.")
        else:
            doc.add_paragraph("[Catchment or sites layer missing – cannot draw site location maps.]")
            _log("WARN", f"Site {sid}: catchment or sites layer missing for location map.")

        site_plot_dir = sites_plots / f"Site_{sid}"
        plot_path = _glob_first(site_plot_dir, cfg.monitoring_plot_patterns) if site_plot_dir.exists() else None
        if plot_path and plot_path.exists():
            fig_counter += 1
            _insert_figure(doc, plot_path, f"Figure {fig_counter}. Monitoring data for Site {sid}.", width_in=6.5)
            _log("OK", f"Added monitoring plot for Site {sid}: {plot_path}")
        else:
            doc.add_paragraph(f"[Monitoring plot not found for Site {sid}. Looked in: {site_plot_dir}]")
            _log("WARN", f"Monitoring plot missing for Site {sid} in: {site_plot_dir}")

        site_data_dir = sites_datasets / f"Site_{sid}"
        site_summary_csv = None
        if site_data_dir.exists():
            matches = list(site_data_dir.glob("*Monitoring*Summary*.csv"))
            if matches:
                site_summary_csv = matches[0]
        if site_summary_csv and site_summary_csv.exists():
            try:
                df_sum = pd.read_csv(site_summary_csv)
                if not df_sum.empty:
                    doc.add_paragraph("")
                    doc.add_heading(f"Monitoring summary for Site {sid}", level=3)
                    cols_site = list(df_sum.columns)
                    tbl = doc.add_table(rows=1, cols=len(cols_site))
                    hdr = tbl.rows[0].cells
                    for i, c in enumerate(cols_site):
                        hdr[i].text = str(c)
                    for _, r in df_sum.iterrows():
                        cells = tbl.add_row().cells
                        for i, c in enumerate(cols_site):
                            val = r.get(c, np.nan)
                            if isinstance(val, (float, np.floating)):
                                cells[i].text = "" if not np.isfinite(val) else f"{val:.3g}"
                            else:
                                cells[i].text = "" if pd.isna(val) else str(val)
                    _log("OK", f"Added monitoring summary table for Site {sid}: {site_summary_csv}")
                else:
                    doc.add_paragraph(f"[Monitoring summary table for Site {sid} is empty.]")
                    _log("WARN", f"Monitoring summary CSV empty for Site {sid}: {site_summary_csv}")
            except Exception as e:
                doc.add_paragraph(f"[Could not add monitoring summary table for Site {sid}: {e}]")
                _log("WARN", f"Could not add monitoring summary table for Site {sid}: {e}")
        else:
            _log("WARN", f"Monitoring summary CSV missing for Site {sid}.")

        if sites_gdf is not None and not sites_gdf.empty and ("Site_id" in sites_gdf.columns):
            row_match = sites_gdf[sites_gdf["Site_id"].astype(str) == str(sid)]
            if not row_match.empty:
                row = row_match.iloc[0]
                le_png = _site_le_panel(cfg, row, sid, site_plot_dir)
                if le_png and le_png.exists():
                    fig_counter += 1
                    _insert_figure(
                        doc, le_png,
                        f"Figure {fig_counter}. NDVI, C-Factor, RUSLE, and SDR-RUSLE ({cfg.le_year_min}–{cfg.le_year_max}) for Site {sid}.",
                        width_in=6.5,
                    )
                    _log("OK", f"Added land/erosion panel for Site {sid}: {le_png}")
                else:
                    _log("WARN", f"Land/erosion panel not generated for Site {sid}.")

        doc.add_heading(f"Additional figures for Site {sid}", level=3)
        site_lu_pct_csv_name = f"Site_{sid}_LU_Percentages.csv"
        site_data_dir = sites_plots / f"Site_{sid}"

        for fname_tpl, cap_tpl in SITE_PLOTS_SEQUENCE:
            filename = fname_tpl.format(sid=sid)
            caption = cap_tpl.format(sid=sid)

            prev_fig_counter = fig_counter
            fig_counter = _insert_site_plot_if_exists(
                doc, site_plot_dir, filename, fig_counter, caption, width_in=6.5
            )

            if (
                prev_fig_counter != fig_counter
                and fname_tpl == "Site_{sid}_LU_Percentages_OverTime.png"
            ):
                lu_pct_csv_path = site_data_dir / site_lu_pct_csv_name
                doc.add_paragraph("")
                doc.add_heading(f"Site {sid} – Land-use percentages table", level=4)
                if lu_pct_csv_path.exists():
                    _add_csv_table(doc, lu_pct_csv_path, round_to=3)
                    _log("OK", f"Added LU percentages table for Site {sid}: {lu_pct_csv_path}")
                else:
                    doc.add_paragraph(f"[{site_lu_pct_csv_name} not found at: {lu_pct_csv_path}]")
                    _log("WARN", f"Missing LU percentages table for Site {sid}: {lu_pct_csv_path}")

        site_daily, dcol_site = _load_site_daily(base, cfg, sid)
        if not site_daily.empty:
            site_annual = _annual_from_daily(site_daily, dcol_site)

            site_hydro_png = _site_hydro_ts_plot(cfg, base, sid, site_annual)
            if site_hydro_png and site_hydro_png.exists():
                fig_counter += 1
                _insert_figure(
                    doc, site_hydro_png,
                    f"Figure {fig_counter}. Annual hydroclimate time series for Site {sid}.",
                    width_in=6.5
                )
                _log("OK", f"Added site hydroclimate plot for Site {sid}: {site_hydro_png}")
            else:
                doc.add_paragraph(f"[Site {sid}: hydroclimate time series not generated.]")
                _log("WARN", f"Site {sid}: hydroclimate time series not generated.")

            doc.add_heading(f"Statistical summary for annual hydroclimate variables – Site {sid}", level=3)
            cols = ["Variable", "Mean", "Median", "Min", "Max", "Std"]
            rows = []
            for var in ["Precipitation", "Mean Temperature", "Runoff", "Actual ET",
                        "Upper Soil Moisture", "Deeper Soil Moisture",
                        "Min Temperature", "Max Temperature"]:
                if isinstance(site_annual, pd.DataFrame) and (var in site_annual.columns):
                    s = pd.to_numeric(site_annual[var], errors="coerce")
                    n = int(s.notna().sum())
                    rows.append({
                        "Variable": var,
                        "Mean":   float(np.nanmean(s)) if n else np.nan,
                        "Median": float(np.nanmedian(s)) if n else np.nan,
                        "Min":    float(np.nanmin(s)) if n else np.nan,
                        "Max":    float(np.nanmax(s)) if n else np.nan,
                        "Std":    float(np.nanstd(s, ddof=1)) if n >= 2 else np.nan,
                    })
            if rows:
                df_site_summary = pd.DataFrame(rows)
                tbl = doc.add_table(rows=1, cols=len(cols))
                hdr = tbl.rows[0].cells
                for i, c in enumerate(cols): hdr[i].text = c
                for _, rsum in df_site_summary.iterrows():
                    rc = tbl.add_row().cells
                    for i, c in enumerate(cols):
                        rc[i].text = _fmt_float(rsum.get(c, np.nan))
                _log("OK", f"Added site hydroclimate summary table for Site {sid}.")
            else:
                doc.add_paragraph(f"[Site {sid}: no annual summary computed (missing inputs).]")
                _log("WARN", f"Site {sid}: no annual hydroclimate summary computed.")

            spi_site_png, sri_site_png, _spi_year, _sri_year = _compute_spi_sri_for_site(cfg, base, sid, site_daily, dcol_site)
            if spi_site_png and spi_site_png.exists():
                fig_counter += 1
                _insert_figure(doc, spi_site_png, f"Figure {fig_counter}. Site {sid} SPI-12 (annual mean of 12-month SPI).")
                _log("OK", f"Added Site {sid} SPI12 figure: {spi_site_png}")
            else:
                doc.add_paragraph(f"[Site {sid}: SPI12 figure not generated.]")
                _log("WARN", f"Site {sid}: SPI12 figure missing.")

            if sri_site_png and sri_site_png.exists():
                fig_counter += 1
                _insert_figure(doc, sri_site_png, f"Figure {fig_counter}. Site {sid} SRI-12 (annual mean of 12-month SRI).")
                _log("OK", f"Added Site {sid} SRI12 figure: {sri_site_png}")
            else:
                doc.add_paragraph(f"[Site {sid}: SRI12 figure not generated.]")
                _log("WARN", f"Site {sid}: SRI12 figure missing.")
        else:
            doc.add_paragraph(f"[Site {sid}: daily hydroclimate CSV not found.]")
            _log("WARN", f"Site {sid}: daily hydroclimate CSV not found.")

    out_docx = base / f"{catchment_name}_Automated_Report.docx"
    try:
        doc.save(out_docx)
        _log("OK", f"Report written: {out_docx}")
    except PermissionError:
        _log("ERR", f"Could not write '{out_docx}'. Please close it in Word and run again.")
    return out_docx